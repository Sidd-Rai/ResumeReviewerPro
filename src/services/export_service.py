"""
Export service for generating reports and improved resumes in PDF/DOCX formats.
"""

from datetime import datetime
from typing import Optional
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.analysis.analysis_result import AnalysisResult
from src.config.settings import COLORS, SCORE_THRESHOLDS


def get_score_color_hex(score: int) -> str:
    """Determine hex color based on score and thresholds."""
    if score >= SCORE_THRESHOLDS["excellent"]:
        return COLORS["excellent"]
    elif score >= SCORE_THRESHOLDS["good"]:
        return COLORS["good"]
    elif score >= SCORE_THRESHOLDS["fair"]:
        return COLORS["fair"]
    else:
        return COLORS["poor"]


def hex_to_rgb(hex_color: str) -> tuple:
    """Convert hex color to RGB tuple for docx."""
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))


class ExportService:
    """Service for exporting analysis results and improved resumes."""
    
    @staticmethod
    def generate_analysis_report_pdf(analysis: AnalysisResult) -> BytesIO:
        """Generate comprehensive PDF report from analysis."""
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor(COLORS["primary"]),
            spaceAfter=12,
            fontName='Helvetica-Bold'
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor(COLORS["primary"]),
            spaceAfter=8,
            spaceBefore=8,
            fontName='Helvetica-Bold'
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=6
        )
        
        story = []
        
        story.append(Paragraph("Resume Analysis Report", title_style))
        story.append(Paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}", normal_style))
        story.append(Spacer(1, 0.2*inch))
        
        story.append(Paragraph("Executive Summary", heading_style))
        story.append(Paragraph(analysis.comprehensive_feedback.executive_summary, normal_style))
        story.append(Spacer(1, 0.15*inch))
        
        story.append(Paragraph("Overall Scores", heading_style))
        
        overall_color = colors.HexColor(get_score_color_hex(analysis.scores.overall))
        ats_color = colors.HexColor(get_score_color_hex(analysis.resume_vs_job.ats_safety_score))
        keyword_color = colors.HexColor(get_score_color_hex(analysis.scores.keyword_density))
        impact_color = colors.HexColor(get_score_color_hex(analysis.scores.impact_quality))
        clarity_color = colors.HexColor(get_score_color_hex(analysis.scores.clarity))
        structure_color = colors.HexColor(get_score_color_hex(analysis.scores.structure))
        
        scores_data = [
            ["Metric", "Score", "Assessment"],
            ["ATS Match", f"{analysis.scores.ats_match}/100", "Keyword alignment with ATS systems"],
            ["Keyword Density", f"{analysis.scores.keyword_density}/100", "Job description keyword coverage"],
            ["Impact Quality", f"{analysis.scores.impact_quality}/100", "Strength of achievements"],
            ["Clarity", f"{analysis.scores.clarity}/100", "Resume readability and structure"],
            ["Structure", f"{analysis.scores.structure}/100", "Overall organization"],
            ["OVERALL", f"{analysis.scores.overall}/100", "Weighted final score"]
        ]
        
        scores_table = Table(scores_data, colWidths=[2*inch, 1.5*inch, 2.5*inch])
        scores_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(COLORS["primary"])),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 11),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (0, 1), ats_color),
            ('BACKGROUND', (0, 2), (0, 2), keyword_color),
            ('BACKGROUND', (0, 3), (0, 3), impact_color),
            ('BACKGROUND', (0, 4), (0, 4), clarity_color),
            ('BACKGROUND', (0, 5), (0, 5), structure_color),
            ('BACKGROUND', (0, 6), (-1, -1), overall_color),
            ('TEXTCOLOR', (0, 6), (-1, -1), colors.whitesmoke),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
            ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
        ]))
        
        story.append(scores_table)
        story.append(Spacer(1, 0.2*inch))
        
        story.append(Paragraph("Job Match Fit", heading_style))
        story.append(Paragraph(
            f"<b>Rating:</b> {analysis.comprehensive_feedback.match_fit.rating}<br/>"
            f"<b>Details:</b> {analysis.comprehensive_feedback.match_fit.explanation}",
            normal_style
        ))
        story.append(Spacer(1, 0.15*inch))
        
        story.append(Paragraph("Top 3 Strengths", heading_style))
        for i, strength in enumerate(analysis.comprehensive_feedback.top_3_strengths, 1):
            story.append(Paragraph(f"{i}. {strength}", normal_style))
        story.append(Spacer(1, 0.15*inch))
        
        story.append(Paragraph("Top 3 Areas for Improvement", heading_style))
        for i, improvement in enumerate(analysis.comprehensive_feedback.top_3_improvements, 1):
            story.append(Paragraph(f"{i}. {improvement}", normal_style))
        story.append(Spacer(1, 0.15*inch))
        
        story.append(Paragraph("Immediate Actions Required", heading_style))
        for action in analysis.comprehensive_feedback.immediate_actions:
            story.append(Paragraph(f"• {action}", normal_style))
        story.append(Spacer(1, 0.15*inch))
        
        if analysis.ats_warnings.critical_fixes:
            story.append(PageBreak())
            story.append(Paragraph("Critical ATS Issues", heading_style))
            for fix in analysis.ats_warnings.critical_fixes:
                story.append(Paragraph(f"⚠️ {fix}", normal_style))
            story.append(Spacer(1, 0.15*inch))
        
        if analysis.resume_vs_job.critical_missing_keywords:
            story.append(Paragraph("Critical Missing Keywords", heading_style))
            keywords_text = ", ".join(analysis.resume_vs_job.critical_missing_keywords[:10])
            story.append(Paragraph(f"{keywords_text}", normal_style))
            story.append(Spacer(1, 0.15*inch))
        
        doc.build(story)
        buffer.seek(0)
        return buffer
    
    @staticmethod
    def generate_improved_resume_docx(
        original_resume_text: str,
        analysis: AnalysisResult,
        improvements_only: bool = True
    ) -> BytesIO:
        """
        Generate improved resume in DOCX format.
        If improvements_only=True, shows suggestions alongside original.
        If False, generates a fully rewritten resume.
        """
        doc = Document()
        
        title = doc.add_heading('Improved Resume', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        metadata = doc.add_paragraph()
        metadata.add_run('Note: ').bold = True
        metadata.add_run(
            'Below are key improvements based on AI analysis. '
            'Customize these suggestions to match your unique experience.'
        )
        
        doc.add_paragraph()
        
        if analysis.rewrite_suggestions.summary_rewrite:
            doc.add_heading('Improved Summary', level=1)
            doc.add_paragraph(analysis.rewrite_suggestions.summary_rewrite)
            doc.add_paragraph()
        
        if analysis.rewrite_suggestions.bullet_improvements:
            doc.add_heading('Improved Bullet Points', level=1)
            
            current_section = None
            for improvement in analysis.rewrite_suggestions.bullet_improvements:
                if improvement.section != current_section:
                    current_section = improvement.section
                    doc.add_heading(f'{current_section} Section', level=2)
                
                p = doc.add_paragraph()
                p.add_run('Improved: ').bold = True
                p.add_run(improvement.improved)
                
                p = doc.add_paragraph()
                p.add_run('Why: ').italic = True
                p.add_run(improvement.reasoning, style='Subtle Emphasis')
                
                doc.add_paragraph()
        
        if analysis.rewrite_suggestions.keywords_to_add:
            doc.add_heading('Keywords to Incorporate', level=1)
            for keyword in analysis.rewrite_suggestions.keywords_to_add:
                doc.add_paragraph(keyword, style='List Bullet')
            doc.add_paragraph()
        
        if analysis.rewrite_suggestions.quick_wins:
            doc.add_heading('Quick Wins (Easy Changes)', level=1)
            for win in analysis.rewrite_suggestions.quick_wins:
                doc.add_paragraph(win, style='List Bullet')
            doc.add_paragraph()
        
        if analysis.rewrite_suggestions.structure_improvements:
            doc.add_heading('Structure Improvements', level=1)
            for improvement in analysis.rewrite_suggestions.structure_improvements:
                doc.add_paragraph(improvement, style='List Bullet')
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    @staticmethod
    def generate_comparison_pdf(
        original_resume_text: str,
        analysis: AnalysisResult
    ) -> BytesIO:
        """Generate side-by-side comparison of before and after."""
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        
        story = []
        
        story.append(Paragraph("<b>Resume Analysis: Before & After</b>", styles['Title']))
        story.append(Spacer(1, 0.2*inch))
        
        story.append(Paragraph("<b>Key Metrics</b>", styles['Heading2']))
        
        overall_color = colors.HexColor(get_score_color_hex(analysis.scores.overall))
        ats_color = colors.HexColor(get_score_color_hex(analysis.resume_vs_job.ats_safety_score))
        keyword_color = colors.HexColor(get_score_color_hex(analysis.scores.keyword_density))
        
        comparison_data = [
            ["Metric", "Current", "Target"],
            ["ATS Score", "Unknown", f"{analysis.scores.ats_match}/100"],
            ["Keyword Match", "Unknown", f"{analysis.scores.keyword_density}%"],
            ["Impact Score", "Unknown", f"{analysis.scores.impact_quality}/10"],
            ["Overall Score", "Unknown", f"{analysis.scores.overall}/100"],
        ]
        
        table = Table(comparison_data, colWidths=[2*inch, 2*inch, 2*inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(COLORS["primary"])),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('BACKGROUND', (2, 1), (2, 1), ats_color),
            ('BACKGROUND', (2, 2), (2, 2), keyword_color),
            ('BACKGROUND', (2, 4), (2, 4), overall_color),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        
        story.append(table)
        story.append(Spacer(1, 0.2*inch))
        
        story.append(Paragraph("<b>Top Improvements to Make</b>", styles['Heading2']))
        for i, imp in enumerate(analysis.comprehensive_feedback.top_3_improvements, 1):
            story.append(Paragraph(f"{i}. {imp}", styles['Normal']))
        
        doc.build(story)
        buffer.seek(0)
        return buffer